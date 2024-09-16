from docx import Document
import random

def extract_metadata(docx, rsid1, rsid2, matching_rsid):
    doc = Document(docx)
    
    # metadata for display
    metadata1 = {
        'title': doc.core_properties.title,
        'created': doc.core_properties.created,
        'paragraphs': []
    }
    metadata2 = {
        'title': doc.core_properties.title,
        'created': doc.core_properties.created,
        'paragraphs': []
    }
    
    dict1, _ = rsid1  # Get only the dictionary part
    dict2, _ = rsid2

    # To track added text for avoiding duplicates
    added_text1 = set()

    # Process dict1 for metadata1
    for rsid, info in dict1.items():
        for text in info['location']:
            if text.strip() and (rsid, text) not in added_text1:
                colour = 'red' if rsid in matching_rsid else random.choice(['blue', 'green', 'yellow', 'cyan', 'grey', 'magenta'])
                metadata1['paragraphs'].append({
                    'rsid': rsid,
                    'colour': colour,
                    'text': text
                })
                added_text1.add((rsid, text))
    
    added_text2 = set()
    # Process dict2 for metadata2
    for rsid, info in dict2.items():
        for text in info['location']:
            if text.strip() and (rsid, text) not in added_text2:
                colour = 'red' if rsid in matching_rsid else random.choice(['blue', 'green', 'yellow', 'cyan', 'grey', 'magenta'])
                metadata2['paragraphs'].append({
                    'rsid': rsid,
                    'colour': colour,
                    'text': text
                })
                added_text2.add((rsid, text))

    return metadata1, metadata2

class DOCX:
    """
    Constants for metadata keys in meta_dict
    """
    CREATED_BY = "Created By"
    DATE_CREATED = "Date Created"
    DATE_LAST_MODIFIED = "Date Last Modified"
    TITLE = "Title"
    VERSION = "Version"

    def __init__(self, document_name):
        self.docx_name = document_name
        self.unique_rsid = {}
        self.paragraphs = {}
        self.num_runs = 0
        self.metadata = {}
        self.properties_dict = {}
        self.settings_rsid = []

    def append_txt(self, paragraph_id, rsid_tag, properties_array, txt):
        """Append text to the appropriate paragraph and process properties."""
        paragraph = self.paragraphs.setdefault(paragraph_id, PARAGRAPH(paragraph_id))
        rsid = self.unique_rsid.setdefault(rsid_tag, RSID(rsid_tag, len(self.unique_rsid)))
        run_index = self.num_runs

        paragraph.append_txt(txt, rsid_tag, rsid.index, run_index)
        self.append_properties(properties_array, run_index)
        self.num_runs += 1

    def append_properties(self, properties_array, run_index):
        for prop in properties_array:
            prop_hash = hash(prop.xml)
            if prop_hash not in self.properties_dict:
                self.properties_dict[prop_hash] = prop
                prop.append_run(run_index, prop.inherit_from)
            else:
                self.properties_dict[prop_hash].append_run(run_index, prop.inherit_from)

    def set_settings_rsid(self, settings_rsid):
        self.settings_rsid = settings_rsid

    def append_metadata(self, key, value):
        self.metadata[key] = value
    
    def get_metadata(self, key):
        return self.metadata.get(key)



class RSID:
    def __init__(self, tag, index):
        self.tag = tag
        self.index = index


class PARAGRAPH:
    def __init__(self, id_):
        self.id = id_
        self.txt_array = []
        self.rsid_array = []
        self.rsid_index_array = []
        self.run_index_array = []

    def append_txt(self, txt, rsid, rsid_index, run_index):
        self.txt_array.append(txt)
        self.rsid_array.append(rsid)
        self.rsid_index_array.append(rsid_index)
        self.run_index_array.append(run_index)

    def get_zip(self):
        return zip(self.txt_array, self.rsid_array, self.rsid_index_array)


class PROPERTY:
    """Inheritance levels"""
    SELF = 0
    PARENT = 1
    GRANDPARENT = 2

    def __init__(self, xml, name, value_dict, inherit_from):
        self.xml = xml
        self.name = name
        self.value_dict = value_dict
        self.inherit_from = inherit_from
        self.runs = []

    def append_run(self, run_index, inherit_value):
        self.runs.append(run_index)

